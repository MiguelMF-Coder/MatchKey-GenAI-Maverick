# app/services/ocr/document_parser.py

import re
import json
import spacy
import unidecode
from pathlib import Path
from typing import Dict, List

from fastapi import UploadFile

# --- Importaciones opcionales según disponibilidad ---
try:
    import docx  # Para Word
except ImportError:
    docx = None
try:
    import pdfplumber  # Para PDF
except ImportError:
    pdfplumber = None
try:
    import pytesseract  # Para OCR
    from PIL import Image
    from pdf2image import convert_from_path
except ImportError:
    pytesseract = None
    Image = None
    convert_from_path = None


# --- CLASE CVREADER ---
class CVReader:
    """Convierte cualquier archivo (PDF, DOCX, IMG) en texto limpio."""

    @staticmethod
    def read_file(file_path: str) -> str:
        path = Path(file_path)
        if not path.exists():
            return ""
        suffix = path.suffix.lower()
        text = ""

        # 1. DOCX (párrafos + tablas)
        if suffix == ".docx" and docx:
            try:
                doc = docx.Document(str(path))
                full_text = []

                # Párrafos normales
                for para in doc.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text)

                # Tablas (muy común en CVs)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                full_text.append(cell.text)

                text = "\n".join(full_text)
            except Exception as e:
                print(f"Error leyendo DOCX: {e}")

        # 2. PDF (texto + fallback OCR si parece escaneado)
        elif suffix == ".pdf" and pdfplumber:
            try:
                with pdfplumber.open(str(path)) as pdf:
                    text = "\n".join([page.extract_text() or "" for page in pdf.pages])

                # Si casi no hay texto, asumimos PDF escaneado y usamos OCR
                if len(text.strip()) < 50 and convert_from_path and pytesseract:
                    images = convert_from_path(str(path))
                    text = "\n".join(
                        [pytesseract.image_to_string(img, lang="spa") for img in images]
                    )
            except Exception as e:
                print(f"Error leyendo PDF: {e}")

        # 3. Imágenes
        elif suffix in [".jpg", ".jpeg", ".png"] and pytesseract and Image:
            try:
                text = pytesseract.image_to_string(Image.open(str(path)), lang="spa")
            except Exception as e:
                print(f"Error leyendo imagen: {e}")

        return CVReader._clean_text(text)

    @staticmethod
    def _clean_text(text: str) -> str:
        if not text:
            return ""
        text = text.replace("\r", "")
        text = re.sub(r"[ \t]{2,}", " ", text)
        return text.strip()


# --- CLASE CVPARSER ---
class CVParser:
    """Coge el texto de CVReader, lo entiende y lo devuelve estructurado."""

    def __init__(self):
        try:
            self.nlp = spacy.load("es_core_news_md")
        except Exception:
            self.nlp = None

        self.SECTION_PATTERNS = {
            "PERFIL": r"(perfil|sobre m[ií]|about me|resumen|intro)",
            "EXPERIENCIA": r"(experiencia|laboral|professional|work|trayectoria|historial)",
            "ESTUDIOS": r"(educaci[oó]n|formaci[oó]n|estudios|acad[eé]mic|t[ií]tulaci[oó]n)",
            "SKILLS": r"(habilidades|skills|competencias|conocimientos|tecnolog[ií]as|herramientas|stack|tech)",
            "IDIOMAS": r"(idiomas|languages)",
            "PROYECTOS": r"(proyectos|projects)",
        }

    def process(self, text: str) -> dict:
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        sections = self._segment_sections(lines)

        # Datos fijos
        name, surname = self._extract_name_strict(lines)
        location = self._extract_location(lines[:20])
        start_date, end_date = self._extract_global_dates(text)

        # Contacto
        contact_info = self._extract_contact_info(text)

        # Secciones
        exp_raw = self._clean_section_headers(
            sections.get("EXPERIENCIA", []) + sections.get("PROYECTOS", []),
            "EXPERIENCIA",
        )
        exp_detail = self._parse_experience_detail(exp_raw)

        estudios_raw = self._clean_section_headers(
            sections.get("ESTUDIOS", []), "ESTUDIOS"
        )

        skills = self._extract_skills_robust(sections.get("SKILLS", []))
        langs = self._extract_languages(sections.get("IDIOMAS", []))

        return {
            "Nombre": name,
            "Apellidos": surname,
            "Ubicacion": location,
            "Contacto": contact_info,
            "Estudios": "\n".join(estudios_raw),
            "Fecha_inicio": start_date,
            "Fecha_finalizacion": end_date,
            "Experiencia": exp_detail,
            "Idiomas": langs,
            "Skills": skills,
            "raw_text_preview": text[:300],
        }

    def _segment_sections(self, lines: List[str]) -> Dict[str, List[str]]:
        segmented = {}
        current = "CONTACTO"
        segmented[current] = []

        for line in lines:
            norm_line = unidecode.unidecode(line.lower())
            matched_section = None
            for key, pat in self.SECTION_PATTERNS.items():
                match_start = re.search(rf"^{pat}\b", norm_line)
                if match_start:
                    if len(line.split()) < 7 or ":" in line[:15]:
                        matched_section = key
                        break

            if matched_section:
                current = matched_section
                if current not in segmented:
                    segmented[current] = []
                segmented[current].append(line)
            else:
                if current not in segmented:
                    segmented[current] = []
                segmented[current].append(line)
        return segmented

    def _extract_contact_info(self, text: str) -> Dict[str, str]:
        contact = {"Email": "", "Telefono": ""}

        # Email
        email_match = re.search(r"[\w\.-]+@[\w\.-]+\.\w+", text)
        if email_match:
            contact["Email"] = email_match.group(0)

        # Teléfono
        phone_pattern = r"(?:\+\d{1,3}[ -]?)?\(?\d{3}\)?[ -]?\d{3}[ -]?\d{3,4}"
        phones = re.findall(phone_pattern, text)
        valid_phones = [p for p in phones if len(re.sub(r"\D", "", p)) >= 9]

        if valid_phones:
            contact["Telefono"] = valid_phones[0].strip()

        return contact

    def _extract_skills_robust(self, lines: List[str]) -> List[str]:
        if not lines:
            return []
        full_text = " ".join(lines)
        pattern_remove = r"\b(Habilidades(\s*técnicas)?|Skills|Competencias|Tecnolog[ií]as|Herramientas|Stack|Tech)\s*:?"
        clean_text = re.sub(pattern_remove, " ", full_text, flags=re.IGNORECASE)
        clean_text = re.sub(r"[()\[\]:]", ",", clean_text)
        tokens = re.split(r"[,;|•·\n]|\s{2,}", clean_text)
        final_skills = []
        blacklist = [
            "avanzado",
            "básico",
            "intermedio",
            "alto",
            "experiencia",
            "nivel",
            "lenguajes",
            "y",
            "e",
        ]
        for t in tokens:
            t = t.strip()
            if 1 < len(t) < 35 and t.lower() not in blacklist:
                final_skills.append(t)
        return list(dict.fromkeys(final_skills))

    def _extract_languages(self, lines: List[str]):
        text = " ".join(lines)
        text = re.sub(r"(Idiomas|Languages)(\s*y\s*otros)?\s*:?", "", text, flags=re.I)
        return [x.strip() for x in re.split(r"[,;\n]", text) if len(x.strip()) > 2]

    def _clean_section_headers(self, lines: List[str], section_key: str) -> List[str]:
        cleaned = []
        pat = self.SECTION_PATTERNS.get(section_key, "")
        for line in lines:
            norm = unidecode.unidecode(line.lower())
            if re.search(rf"^{pat}.{{0,5}}$", norm) and len(line.split()) < 4:
                continue
            cleaned.append(line)
        return cleaned

    def _extract_name_strict(self, lines):
        if not lines:
            return "", ""
        first = lines[0].strip()
        if re.search(r"[@\d]", first):
            return "", ""
        parts = first.split()
        if 2 <= len(parts) <= 4:
            return parts[0], " ".join(parts[1:])
        return "", ""

    def _extract_location(self, lines):
        text = "\n".join(lines)
        m = re.search(r"([A-ZÁÉÍÓÚ][a-zñáéíóú]+),\s*([A-ZÁÉÍÓÚ][a-zñáéíóú]+)", text)
        return m.group(0) if m else ""

    def _extract_global_dates(self, text):
        years = sorted(
            [int(y) for y in re.findall(r"\b(19\d{2}|20[0-3]\d)\b", text)]
        )
        if not years:
            return "", ""
        start = str(years[0])
        end = (
            "Actualidad"
            if re.search(r"\b(actualidad|presente)\b", text, re.I)
            else str(years[-1])
        )
        return start, end

    def _parse_experience_detail(self, lines: List[str]):
        res = {"Empleos": [], "Practicas": [], "Proyectos": []}
        full = "\n".join(lines)
        blocks = re.split(r"\n\n+", full)
        for b in blocks:
            b = b.strip()
            if not b:
                continue
            lb = b.lower()
            if "prácticas" in lb or "practicas" in lb or "beca" in lb:
                res["Practicas"].append(b)
            elif "proyecto" in lb:
                res["Proyectos"].append(b)
            else:
                res["Empleos"].append(b)
        return res


# --- FUNCIÓN PÚBLICA PARA EL ENDPOINT ---

import tempfile
import os


def parse_cv_upload(upload_file: UploadFile) -> dict:
    """
    Guarda temporalmente el CV subido, lanza CVReader + CVParser
    y devuelve un dict con la info estructurada.
    """
    suffix = ""
    if upload_file.filename and "." in upload_file.filename:
        suffix = "." + upload_file.filename.split(".")[-1]

    # Guardamos en un fichero temporal
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(upload_file.file.read())
        tmp_path = tmp.name

    try:
        reader = CVReader()
        text = reader.read_file(tmp_path)
        if not text:
            return {"error": "No se pudo extraer texto del CV."}

        parser = CVParser()
        data = parser.process(text)
        return data
    finally:
        # Limpiamos el fichero temporal
        try:
            os.remove(tmp_path)
        except OSError:
            pass
